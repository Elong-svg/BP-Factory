#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word Pro - 专业级一键转换脚本
应用 minimax-docx 企业级设计系统

用法：
    python markdown_to_word_pro.py input.md [output.docx]
"""

import sys
import os
import markdown
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from bs4 import BeautifulSoup
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import xml.etree.ElementTree as ET


# ============================================================================
# 模板配置
# ============================================================================

class TemplateConfig:
    """模板配置类 - 从 XML 文件加载样式配置"""
    
    def __init__(self, template_name='user_default'):
        """
        初始化模板配置
        
        Args:
            template_name: 模板名称 
                - 'user_default': 用户自定义默认（黑体标题 + 宋体正文 + 楷体图注）
                - 'corporate': 企业报告模板
                - 'academic': 学术论文模板
                - 'default': 通用文档模板
        """
        self.template_name = template_name
        self.config = self._load_template(template_name)
        
        # 从配置中提取各项参数
        self._extract_config()
    
    def _load_template(self, template_name):
        """加载模板配置文件"""
        # 获取脚本所在目录
        script_dir = os.path.dirname(os.path.abspath(__file__))
        template_file = os.path.join(script_dir, '..', 'assets', 'styles', f'{template_name}_styles.xml')
        
        # 如果文件不存在，使用默认模板
        if not os.path.exists(template_file):
            print(f'[WARNING] 模板文件不存在：{template_file}，使用默认模板 corporate')
            template_file = os.path.join(script_dir, '..', 'assets', 'styles', 'corporate_styles.xml')
        
        print(f'[INFO] 使用模板：{template_name}')
        tree = ET.parse(template_file)
        return tree.getroot()
    
    def _extract_config(self):
        """从 XML 配置中提取各项参数"""
        # 颜色系统（带安全默认值）
        colors = self.config.find('colors')
        self.COLOR_PRIMARY = self._parse_color(colors.find('color[@name="primary"]')) or RGBColor(0x1A, 0x3A, 0x6B)
        self.COLOR_SECONDARY = self._parse_color(colors.find('color[@name="secondary"]')) or RGBColor(0x2B, 0x57, 0x9A)
        self.COLOR_BODY = self._parse_color(colors.find('color[@name="body"]')) or RGBColor(0x2D, 0x2D, 0x2D)
        self.COLOR_TABLE_HEADER_BG = self._parse_color(colors.find('color[@name="table_header"]') or colors.find('color[@name="secondary"]')) or RGBColor(0x2B, 0x57, 0x9A)
        self.COLOR_TABLE_STRIPE = self._parse_color(colors.find('color[@name="table_stripe"]')) or RGBColor(0xE8, 0xEE, 0xF4)
        self.COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
        
        # 字体配置（支持更细粒度的控制）
        fonts = self.config.find('fonts')
        
        # 英文字体（备用）
        self.FONT_EN_TITLE = fonts.find('font_family[@type="title_en"]').text if fonts.find('font_family[@type="title_en"]') is not None else 'Calibri Light'
        self.FONT_EN_BODY = fonts.find('font_family[@type="body_en"]').text if fonts.find('font_family[@type="body_en"]') is not None else 'Calibri'
        
        # 中文字体 - 支持不同级别使用不同字体
        self.FONT_CN_H1 = fonts.find('font_family[@type="title_h1"]').text if fonts.find('font_family[@type="title_h1"]') is not None else fonts.find('font_family[@type="title_cn"]').text
        self.FONT_CN_H2 = fonts.find('font_family[@type="title_h2"]').text if fonts.find('font_family[@type="title_h2"]') is not None else fonts.find('font_family[@type="title_cn"]').text
        self.FONT_CN_H3 = fonts.find('font_family[@type="title_h3"]').text if fonts.find('font_family[@type="title_h3"]') is not None else fonts.find('font_family[@type="title_cn"]').text
        self.FONT_CN_BODY = fonts.find('font_family[@type="body"]').text if fonts.find('font_family[@type="body"]') is not None else fonts.find('font_family[@type="body_cn"]').text
        self.FONT_CN_CAPTION = fonts.find('font_family[@type="caption"]').text if fonts.find('font_family[@type="caption"]') is not None else fonts.find('font_family[@type="body_cn"]').text
        
        # 字号配置
        font_sizes = self.config.find('font_sizes')
        self.FONT_SIZE_H1 = Pt(float(font_sizes.find('size[@element="h1"]').get('points')))
        self.FONT_SIZE_H2 = Pt(float(font_sizes.find('size[@element="h2"]').get('points')))
        self.FONT_SIZE_H3 = Pt(float(font_sizes.find('size[@element="h3"]').get('points')))
        self.FONT_SIZE_BODY = Pt(float(font_sizes.find('size[@element="body"]').get('points')))
        self.FONT_SIZE_CAPTION = Pt(float(font_sizes.find('size[@element="caption"]').get('points'))) if font_sizes.find('size[@element="caption"]') is not None else Pt(10.5)
        
        # 间距配置
        spacing = self.config.find('spacing')
        self.SPACING_H1_BEFORE = Pt(int(spacing.find('heading_1').get('before')))
        self.SPACING_H1_AFTER = Pt(int(spacing.find('heading_1').get('after')))
        self.SPACING_H2_BEFORE = Pt(int(spacing.find('heading_2').get('before')))
        self.SPACING_H2_AFTER = Pt(int(spacing.find('heading_2').get('after')))
        self.SPACING_H3_BEFORE = Pt(int(spacing.find('heading_3').get('before')))
        self.SPACING_H3_AFTER = Pt(int(spacing.find('heading_3').get('after')))
        self.SPACING_BODY_AFTER = Pt(int(spacing.find('body').get('after')))
        
        # 行距
        self.LINE_SPACING = float(spacing.find('line_spacing').get('multiplier'))
        
        # 页面布局
        page_layout = self.config.find('page_layout')
        paper = page_layout.find('paper_size')
        unit = paper.get('unit')
        if unit == 'mm':
            self.PAGE_WIDTH = Cm(float(paper.get('width')) / 10)
            self.PAGE_HEIGHT = Cm(float(paper.get('height')) / 10)
        elif unit == 'inch':
            self.PAGE_WIDTH = Cm(float(paper.get('width')) * 2.54)
            self.PAGE_HEIGHT = Cm(float(paper.get('height')) * 2.54)
        
        margins = page_layout.find('margins')
        for margin in margins.findall('margin'):
            pos = margin.get('position')
            value = float(margin.get('value'))
            unit = margin.get('unit')
            cm_value = value * 2.54 if unit == 'inch' else value
            
            if pos == 'top':
                self.MARGIN_TOP = Cm(cm_value)
            elif pos == 'bottom':
                self.MARGIN_BOTTOM = Cm(cm_value)
            elif pos == 'left':
                self.MARGIN_LEFT = Cm(cm_value)
            elif pos == 'right':
                self.MARGIN_RIGHT = Cm(cm_value)
    
    def _parse_color(self, color_elem):
        """解析颜色元素，安全 fallback（不再返回纯黑）"""
        if color_elem is None:
            return None  # 返回 None 而非黑色，让调用方决定
        
        hex_str = color_elem.get('hex', '').lstrip('#')
        if not hex_str or len(hex_str) != 6:
            return None
        
        return RGBColor(
            int(hex_str[0:2], 16),
            int(hex_str[2:4], 16),
            int(hex_str[4:6], 16)
        )


# ============================================================================
# 样式应用函数
# ============================================================================

def set_chinese_font(run, font_name):
    """设置中文字体"""
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)


def apply_heading1_style(paragraph, text, config):
    """应用一级标题样式（H1）"""
    run = paragraph.add_run(text)
    run.font.name = config.FONT_EN_TITLE
    run.font.size = config.FONT_SIZE_H1
    run.font.bold = True
    run.font.color.rgb = config.COLOR_PRIMARY
    set_chinese_font(run, config.FONT_CN_H1)  # H1 使用独立字体
    
    paragraph.paragraph_format.space_before = config.SPACING_H1_BEFORE
    paragraph.paragraph_format.space_after = config.SPACING_H1_AFTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = config.LINE_SPACING


def apply_heading2_style(paragraph, text, config):
    """应用二级标题样式（H2）"""
    run = paragraph.add_run(text)
    run.font.name = config.FONT_EN_TITLE
    run.font.size = config.FONT_SIZE_H2
    run.font.bold = True
    run.font.color.rgb = config.COLOR_PRIMARY
    set_chinese_font(run, config.FONT_CN_H2)  # H2 使用独立字体
    
    paragraph.paragraph_format.space_before = config.SPACING_H2_BEFORE
    paragraph.paragraph_format.space_after = config.SPACING_H2_AFTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = config.LINE_SPACING


def apply_heading3_style(paragraph, text, config):
    """应用三级标题样式（H3）"""
    run = paragraph.add_run(text)
    run.font.name = config.FONT_EN_TITLE
    run.font.size = config.FONT_SIZE_H3
    run.font.bold = True
    run.font.color.rgb = config.COLOR_PRIMARY
    set_chinese_font(run, config.FONT_CN_H3)  # H3 使用独立字体（宋体）
    
    paragraph.paragraph_format.space_before = config.SPACING_H3_BEFORE
    paragraph.paragraph_format.space_after = config.SPACING_H3_AFTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = config.LINE_SPACING


def apply_body_style(paragraph, text, config):
    """应用正文样式"""
    run = paragraph.add_run(text)
    run.font.name = config.FONT_EN_BODY
    run.font.size = config.FONT_SIZE_BODY
    run.font.color.rgb = config.COLOR_BODY
    set_chinese_font(run, config.FONT_CN_BODY)
    
    paragraph.paragraph_format.space_after = config.SPACING_BODY_AFTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = config.LINE_SPACING


def create_table_with_style(doc, headers, rows, config):
    """创建带样式的表格（三线表）"""
    # 创建表格 - 不使用预设样式，完全自定义
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    
    # 设置表格边框（细灰线）
    tbl = table._element
    
    # 获取或创建 tblPr
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # 创建边框
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    
    # 设置表头
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        
        # 设置表头样式：蓝色背景 + 白色粗体文字
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 清除原有 runs，重新添加
        for run in paragraph.runs:
            run.text = ''
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run('')
        run.text = header_text
        run.font.bold = True
        run.font.color.rgb = config.COLOR_WHITE
        run.font.size = config.FONT_SIZE_BODY
        set_chinese_font(run, config.FONT_CN_BODY)
        
        # 设置单元格背景色（使用 shading）
        tc = cell._element.tcPr
        tcShd = OxmlElement('w:shd')
        tcShd.set(qn('w:val'), 'clear')
        tcShd.set(qn('w:color'), 'auto')
        # Word 的 shading fill 需要 RGB 十六进制
        header_color_hex = '{:02X}{:02X}{:02X}'.format(
            config.COLOR_TABLE_HEADER_BG[0] if isinstance(config.COLOR_TABLE_HEADER_BG, tuple) else 0x2B,
            config.COLOR_TABLE_HEADER_BG[1] if isinstance(config.COLOR_TABLE_HEADER_BG, tuple) else 0x57,
            config.COLOR_TABLE_HEADER_BG[2] if isinstance(config.COLOR_TABLE_HEADER_BG, tuple) else 0x9A,
        )
        # RGBColor 是特殊类型，直接转 hex
        try:
            header_color_hex = str(config.COLOR_TABLE_HEADER_BG)
        except:
            header_color_hex = '2B579A'
        tcShd.set(qn('w:fill'), header_color_hex)
        tc.append(tcShd)
    
    # 设置数据行
    for row_idx, row_data in enumerate(rows, start=1):
        row = table.rows[row_idx]
        for i, cell_text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = cell_text
            
            paragraph = cell.paragraphs[0]
            if paragraph.runs:
                run = paragraph.runs[0]
            else:
                run = paragraph.add_run(cell_text)
            run.font.size = config.FONT_SIZE_BODY
            set_chinese_font(run, config.FONT_CN_BODY)
            
            # 隔行条纹（使用 shading）
            if row_idx % 2 == 0 and config.COLOR_TABLE_STRIPE:
                tc = cell._element.tcPr
                tcShd = OxmlElement('w:shd')
                tcShd.set(qn('w:val'), 'clear')
                tcShd.set(qn('w:color'), 'auto')
                try:
                    stripe_color_hex = str(config.COLOR_TABLE_STRIPE)
                except:
                    stripe_color_hex = 'E8EEF4'
                if len(stripe_color_hex) == 6:
                    tcShd.set(qn('w:fill'), stripe_color_hex)
                tc.append(tcShd)
    
    return table


# ============================================================================
# 核心转换函数
# ============================================================================

def markdown_to_word(md_file, docx_file, config):
    """
    将 Markdown 文件转换为 Word 文档
    
    Args:
        md_file: 输入 Markdown 文件路径
        docx_file: 输出 Word 文件路径
        config: 模板配置对象
    """
    print(f'[INFO] 正在读取 Markdown 文件：{md_file}')
    
    # 1. 读取 Markdown 文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_text = f.read()
    
    print(f'[INFO] Markdown 文件大小：{len(md_text)} 字符')
    
    # 1.5 预处理 Markdown：确保表格前后有空行
    import re
    # 匹配 Markdown 表格（包含表头、分隔线、数据行）
    table_pattern = r'^(\s*\|.+\|\s*\n\s*\|[-|\s]+\|\s*\n(?:\s*\|.+\|\s*\n)*)'
    
    def ensure_blank_lines(match):
        """确保表格前后有空行"""
        table = match.group(1)
        return '\n' + table + '\n'
    
    md_text = re.sub(table_pattern, ensure_blank_lines, md_text, flags=re.MULTILINE)
    
    # 2. 解析 Markdown → HTML
    print('[INFO] 正在解析 Markdown...')
    html = markdown.markdown(md_text, extensions=['extra', 'codehilite', 'tables', 'toc'])
    
    # 3. 解析 HTML
    print('[INFO] 正在解析 HTML 结构...')
    soup = BeautifulSoup(html, 'html.parser')
    
    # 4. 创建 Word 文档
    print('[INFO] 正在创建 Word 文档...')
    doc = Document()
    
    # 5. 设置页面布局
    section = doc.sections[0]
    section.page_width = config.PAGE_WIDTH
    section.page_height = config.PAGE_HEIGHT
    section.top_margin = config.MARGIN_TOP
    section.bottom_margin = config.MARGIN_BOTTOM
    section.left_margin = config.MARGIN_LEFT
    section.right_margin = config.MARGIN_RIGHT
    
    # 6. 遍历 HTML 元素并转换为 Word
    print('[INFO] 正在转换内容...')
    
    for element in soup.children:
        # 跳过空白元素
        if not hasattr(element, 'name') or element.name is None:
            continue
        
        # 处理标题
        if element.name == 'h1':
            p = doc.add_paragraph()
            apply_heading1_style(p, element.get_text(), config)
        
        elif element.name == 'h2':
            p = doc.add_paragraph()
            apply_heading2_style(p, element.get_text(), config)
        
        elif element.name == 'h3':
            p = doc.add_paragraph()
            apply_heading3_style(p, element.get_text(), config)
        
        elif element.name == 'h4':
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.font.name = config.FONT_EN_TITLE
            h4_size = Pt(13) if config.FONT_SIZE_H1 > Pt(20) else Pt(12)
            run.font.size = h4_size
            run.font.bold = True
            run.font.color.rgb = config.COLOR_SECONDARY
            set_chinese_font(run, config.FONT_CN_H3)
            p.paragraph_format.space_after = config.SPACING_BODY_AFTER
        
        elif element.name == 'p':
            # 处理段落（包含粗体、斜体等）
            p = doc.add_paragraph()
            
            # 遍历段落内的所有元素
            for child in element.children:
                if hasattr(child, 'name'):
                    if child.name == 'strong' or child.name == 'b':
                        run = p.add_run(child.get_text())
                        run.font.bold = True
                    elif child.name == 'em' or child.name == 'i':
                        run = p.add_run(child.get_text())
                        run.font.italic = True
                    elif child.name == 'code':
                        run = p.add_run(child.get_text())
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                    else:
                        run = p.add_run(child.get_text())
                elif hasattr(child, 'string'):
                    run = p.add_run(str(child))
            
            apply_body_style(p, '', config)  # 应用正文样式
        
        elif element.name == 'ul':
            # 无序列表
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                
                # 处理列表项内容
                for child in li.children:
                    if hasattr(child, 'name'):
                        if child.name == 'strong' or child.name == 'b':
                            run = p.add_run(child.get_text())
                            run.font.bold = True
                        elif child.name == 'em' or child.name == 'i':
                            run = p.add_run(child.get_text())
                            run.font.italic = True
                        else:
                            run = p.add_run(child.get_text())
                    elif hasattr(child, 'string'):
                        run = p.add_run(str(child))
                
                p.paragraph_format.space_after = config.SPACING_BODY_AFTER
        
        elif element.name == 'ol':
            # 有序列表
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph()
                p.style = 'List Number'
                
                # 处理列表项内容
                for child in li.children:
                    if hasattr(child, 'name'):
                        if child.name == 'strong' or child.name == 'b':
                            run = p.add_run(child.get_text())
                            run.font.bold = True
                        elif child.name == 'em' or child.name == 'i':
                            run = p.add_run(child.get_text())
                            run.font.italic = True
                        else:
                            run = p.add_run(child.get_text())
                    elif hasattr(child, 'string'):
                        run = p.add_run(str(child))
                
        
        elif element.name == 'table':
            # 表格
            headers = []
            rows = []
            
            for tr in element.find_all('tr'):
                row_data = [td.get_text().strip() for td in tr.find_all(['th', 'td'])]
                if tr.find('th'):
                    headers = row_data
                else:
                    rows.append(row_data)
            
            if headers:
                print(f'  [TABLE] 创建表格：{len(headers)}列 × {len(rows)+1}行')
                create_table_with_style(doc, headers, rows, config)
        
        elif element.name == 'blockquote':
            # 引用块
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.right_indent = Cm(1.27)
            run = p.add_run(element.get_text())
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            set_chinese_font(run, config.FONT_CN_BODY)
        
        elif element.name == 'hr':
            # 分页符
            doc.add_page_break()
    
    # 7. 保存 Word 文档
    print(f'[INFO] 正在保存 Word 文档：{docx_file}')
    doc.save(docx_file)
    
    print(f'[SUCCESS] 转换完成！')
    print(f'   输入：{md_file}')
    print(f'   输出：{docx_file}')


# ============================================================================
# 主程序
# ============================================================================

def main():
    """主程序入口"""
    import argparse
    
    # 解析命令行参数
    parser = argparse.ArgumentParser(
        description='Markdown to Word Pro - 专业级一键转换工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  python markdown_to_word_pro.py 输入.md                          # 使用默认模板（用户自定义）
  python markdown_to_word_pro.py 输入.md -t user_default          # 使用用户自定义默认模板（黑体标题 + 宋体正文）
  python markdown_to_word_pro.py 输入.md -t corporate             # 使用企业报告模板
  python markdown_to_word_pro.py 输入.md -t academic              # 使用学术论文模板
  python markdown_to_word_pro.py 输入.md -t default               # 使用通用文档模板
  python markdown_to_word_pro.py 输入.md -o 输出.docx             # 指定输出文件
  python markdown_to_word_pro.py 输入.md -t academic -o 论文.docx # 组合使用
        '''
    )
    
    parser.add_argument('input', help='输入 Markdown 文件路径')
    parser.add_argument('-o', '--output', help='输出 Word 文件路径（默认：[输入文件名]_专业版.docx）')
    parser.add_argument('-t', '--template', 
                       choices=['user_default', 'corporate', 'academic', 'default'],
                       default='user_default',  # v2.0: 已重写，专业中文文档首选
                       help='模板类型：user_default（中文专业，默认推荐）, corporate（企业报告，英文风）, academic（学术论文）, default（通用）')
    
    args = parser.parse_args()
    
    md_file = args.input
    
    # 确定输出文件路径
    if args.output:
        docx_file = args.output
    else:
        # 默认输出文件名：[原名]_专业版.docx
        base_name = os.path.splitext(md_file)[0]
        docx_file = f'{base_name}_专业版.docx'
    
    # 检查输入文件是否存在
    if not os.path.exists(md_file):
        print(f'[ERROR] 错误：文件不存在 - {md_file}')
        sys.exit(1)
    
    try:
        # 加载模板配置
        config = TemplateConfig(args.template)
        
        # 执行转换
        markdown_to_word(md_file, docx_file, config)
        print('')
        print('[SUCCESS] 转换成功！')
        print(f'  模板：{args.template}')
        print(f'  输出：{docx_file}')
    except Exception as e:
        print(f'[ERROR] 转换失败：{str(e)}')
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
